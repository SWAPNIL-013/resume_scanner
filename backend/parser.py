import fitz  
from docx import Document
import os

def extract_text_and_links(file_path: str):
    """
    Extracts text + embedded hyperlinks from PDF, DOCX, or TXT files.
    Skips 'mailto:' (emails) and 'tel:' (phone numbers).

    Args:
        file_path (str): Path to resume file

    Returns:
        tuple(str, list): (extracted text, list of urls)
    """
    if not os.path.exists(file_path):
        print(f"⚠️ File does not exist: {file_path}")
        return "", []

    ext = os.path.splitext(file_path)[1].lower()

    text = ""
    links = []

    try:
        if ext == ".pdf":
            with fitz.open(file_path) as pdf:
                for page in pdf:
                    text += page.get_text("text")
                    # extract embedded links
                    for link in page.get_links():
                        if "uri" in link:
                            uri = link["uri"]
                            # Skip email and phone links
                            if not (uri.lower().startswith("mailto:") or uri.lower().startswith("tel:")):
                                links.append(uri)

        elif ext == ".docx":
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            # extract hyperlinks
            rels = doc.part.rels
            for rel in rels.values():
                if "hyperlink" in rel.reltype:
                    uri = rel.target_ref
                    # Skip email and phone links
                    if not (uri.lower().startswith("mailto:") or uri.lower().startswith("tel:")):
                        links.append(uri)

        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

        else:
            print(f"⚠️ Unsupported file format: {file_path}")
            return "", []

    except Exception as e:
        print(f"⚠️ Failed to extract text/links from {file_path}: {e}")
        return "", []

    return text.strip(), links


# ----------------- Test Runner -----------------
if __name__ == "__main__":
    # Path to resume file
    resume_file = "samples/Sakshi_Kusmude_Resume (4).pdf"

    resume_text, resume_links = extract_text_and_links(resume_file)

    print("----- Extracted Resume Text (first 500 chars) -----")
    print(resume_text[:])

    print("\n----- Extracted Links -----")
    if resume_links:
        for link in resume_links:
            print(link)
    else:
        print("No links found")

def extract_text_from_jd(file_path: str):
    """
    Extracts text from pdf, docx files

    Args:
        file_path (str): Path to resume file

    Returns:
        str: extracted text
    """
    if not os.path.exists(file_path):
        print(f"⚠️ File does not exist: {file_path}")
        return ""

    ext = os.path.splitext(file_path)[1].lower()

    text = ""
  

    try:
        if ext == ".pdf":
            with fitz.open(file_path) as pdf:
                for page in pdf:
                    text += page.get_text("text")

        elif ext == ".docx":
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

        else:
            print(f"⚠️ Unsupported file format: {file_path}")
            return ""

    except Exception as e:
        print(f"⚠️ Failed to extract text from {file_path}: {e}")
        return ""

    return text.strip()